from flask import Flask, request, send_file, jsonify, send_from_directory, after_this_request
from flask_cors import CORS
import os
import tempfile
from pdf2docx import Converter as Pdf2DocxConverter
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt
import fitz
import cv2
import numpy as np
from PIL import Image

# --- ★★★ पहला बदलाव: लोकल पाथ को हटा दें ★★★ ---
# Render पर इनकी ज़रूरत नहीं है
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# poppler_path_for_code = r"C:\poppler-24.08.0\Library\bin"

app = Flask(__name__)
CORS(app)

# --- helper function ---
def is_hindi_or_symbol(char):
    if '\u0900' <= char <= '\u097F':
        return True
    if char in ".,!?;:/\\()[]{}'\"-–—…%₹$&*@#=+|_^~<>`":
        return True
    if char == " ":
        return True
    return False

def is_pdf_scanned(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        if doc.page_count == 0:
            doc.close()
            return True
        for page in doc:
            if page.get_text("text").strip():
                doc.close()
                return False
        doc.close()
        return True
    except:
        return True

def process_image_to_docx(image_path, output_path):
    doc = Document()
    pil_image = Image.open(image_path)
    open_cv_image = np.array(pil_image)[:, :, ::-1].copy()

    gray_image = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
    _, processed_image = cv2.threshold(gray_image, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)

    text = pytesseract.image_to_string(processed_image, lang='hin+eng')
    p = doc.add_paragraph()

    current_run_text = ""
    is_current_run_hindi_flag = False
    if text:
        is_current_run_hindi_flag = is_hindi_or_symbol(text[0])

    for char in text:
        if is_hindi_or_symbol(char) == is_current_run_hindi_flag:
            current_run_text += char
        else:
            run = p.add_run(current_run_text)
            run.font.name = 'Mangal' if is_current_run_hindi_flag else 'Arial'
            run.font.size = Pt(11)
            current_run_text = char
            is_current_run_hindi_flag = not is_current_run_hindi_flag

    if current_run_text:
        run = p.add_run(current_run_text)
        run.font.name = 'Mangal' if is_current_run_hindi_flag else 'Arial'
        run.font.size = Pt(11)

    doc.save(output_path)

@app.route('/')
def index():
    # --- ★★★ दूसरा बदलाव: Health Check के लिए ★★★ ---
    # इसे send_from_directory से बदलकर एक सरल संदेश भेजें
    return "PDF/Image Converter Backend is running!", 200

@app.route('/convert', methods=['POST'])
def convert_pdf_or_image_to_word():
    if 'file' not in request.files:
        return jsonify({'error': 'कोई फ़ाइल प्राप्त नहीं हुई।'}), 400
    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({'error': 'कोई फ़ाइल चयनित नहीं की गई।'}), 400

    input_temp_file = tempfile.NamedTemporaryFile(delete=False)
    input_path = input_temp_file.name
    uploaded_file.save(input_path)
    input_temp_file.close()

    output_temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    output_path = output_temp_file.name
    output_temp_file.close()

    try:
        file_ext = os.path.splitext(uploaded_file.filename)[1].lower()

        if file_ext == ".pdf":
            scanned = is_pdf_scanned(input_path)
            if not scanned:
                print("नेटिव PDF मिला। pdf2docx का उपयोग...")
                cv = Pdf2DocxConverter(input_path)
                cv.convert(output_path)
                cv.close()
            else:
                print("स्कैन PDF मिला। OCR का उपयोग...")
                # --- ★★★ तीसरा बदलाव: poppler_path को हटा दें ★★★ ---
                images = convert_from_path(input_path, dpi=300)
                doc = Document()
                for i, pil_image in enumerate(images):
                    open_cv_image = np.array(pil_image)[:, :, ::-1].copy()
                    gray_image = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
                    _, processed_image = cv2.threshold(gray_image, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
                    text = pytesseract.image_to_string(processed_image, lang='hin+eng')
                    p = doc.add_paragraph()

                    current_run_text = ""
                    is_current_run_hindi_flag = False
                    if text:
                        is_current_run_hindi_flag = is_hindi_or_symbol(text[0])

                    for char in text:
                        if is_hindi_or_symbol(char) == is_current_run_hindi_flag:
                            current_run_text += char
                        else:
                            run = p.add_run(current_run_text)
                            run.font.name = 'Mangal' if is_current_run_hindi_flag else 'Arial'
                            run.font.size = Pt(11)
                            current_run_text = char
                            is_current_run_hindi_flag = not is_current_run_hindi_flag

                    if current_run_text:
                        run = p.add_run(current_run_text)
                        run.font.name = 'Mangal' if is_current_run_hindi_flag else 'Arial'
                        run.font.size = Pt(11)

                    if i < len(images) - 1:
                        doc.add_page_break()
                doc.save(output_path)

        else:
            print("Image file मिला। OCR का उपयोग...")
            process_image_to_docx(input_path, output_path)

        @after_this_request
        def cleanup(response):
            try:
                os.remove(input_path)
                os.remove(output_path)
            except:
                pass
            return response

        return send_file(
            output_path,
            as_attachment=True,
            download_name=os.path.splitext(uploaded_file.filename)[0] + '.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'कन्वर्ज़न में समस्या आई।'}), 500

# --- Gunicorn इस हिस्से का उपयोग नहीं करता, लेकिन इसे रखने में कोई हर्ज नहीं ---
if __name__ == '__main__':
    app.run(debug=True, port=5000)
